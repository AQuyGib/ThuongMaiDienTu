from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


INPUT_PATH = Path(r"G:\ThuongMaiDienTu\Nhom_g_bao_cao_lap_trinh_bach_end_web2_chinh_sua_Xuan_Hoa_bang_tong_quan_da_chinh.docx")
OUTPUT_PATH = Path(r"G:\ThuongMaiDienTu\Nhom_g_bao_cao_lap_trinh_bach_end_web2_chinh_sua_Xuan_Hoa_phan_chia_chi_tiet_da_chinh.docx")


ROWS = [
    {
        "stt": "1",
        "member": "Nguyễn Anh Quý\nPhân hệ trải nghiệm khách hàng, AI & Marketing",
        "features": [
            "AI Chatbot Gemini RAG và tạo phiếu sửa chữa từ hội thoại.",
            "CRUD bài viết/blog công nghệ và kiểm duyệt nội dung.",
            "CRUD khách hàng, quản lý CRM và trạng thái người dùng.",
            "Flash Sale giờ vàng, giới hạn số lượng và chống overselling.",
            "Gợi ý bán chéo, combo bundle và FBT suggestion.",
            "Bộ lọc sản phẩm nâng cao theo danh mục/thông số.",
            "So sánh sản phẩm và đồng bộ danh sách so sánh.",
            "Tích điểm thành viên, hạng VIP và lịch sử đổi thưởng.",
            "Thông báo hệ thống/chiến dịch thông báo.",
            "Trang chủ, banner, section sản phẩm và tối ưu UI.",
        ],
        "files": [
            "app/Http/Controllers/ChatbotController.php",
            "app/Services/CrossSellService.php; app/Services/FlashSaleService.php",
            "app/Http/Controllers/ProductFilterController.php; app/Http/Controllers/CompareController.php",
            "app/Http/Controllers/RewardsController.php; app/Http/Controllers/Admin/RewardsController.php",
            "app/Http/Controllers/Admin/NotificationCampaignController.php; app/Services/NotificationService.php",
            "app/Http/Controllers/ArticleFrontendController.php; app/Http/Controllers/Admin/ArticleController.php",
            "app/Http/Controllers/Admin/CustomerController.php; app/Http/Controllers/Admin/HomeSectionController.php",
            "resources/views/partials/chatbot.blade.php; resources/views/frontend/products/*.blade.php",
        ],
        "result": "Hoàn thành nhóm chức năng tương tác khách hàng, marketing, AI tư vấn và các module tăng chuyển đổi bán hàng.",
    },
    {
        "stt": "2",
        "member": "Huỳnh Văn Vĩnh Em\nPhân hệ giỏ hàng, thanh toán, vận chuyển",
        "features": [
            "Quản lý giỏ hàng AJAX/session: thêm, cập nhật, xóa, chọn sản phẩm.",
            "Thanh toán đơn hàng, tạo order và cập nhật trạng thái thanh toán.",
            "Sinh mã QR chuyển khoản phục vụ thanh toán nhanh.",
            "Tính phí vận chuyển theo thông tin giao hàng.",
            "Mã giảm giá/voucher: validate điều kiện, lượt dùng và giá trị giảm.",
            "Tra cứu đơn hàng và theo dõi trạng thái xử lý.",
            "In hóa đơn/bill cho đơn hàng.",
            "Đăng ký trả góp phía khách hàng.",
            "Hủy đơn, timeout đơn, trả tồn kho khi giao dịch không hoàn tất.",
            "Luồng bán hàng tại quầy/POS cơ bản dựa trên giỏ hàng và in bill.",
        ],
        "files": [
            "app/Http/Controllers/Admin/CartController.php",
            "app/Services/CartService.php; app/Services/FlashSaleService.php; app/Services/PointsService.php",
            "app/Http/Controllers/InstallmentController.php; app/Http/Controllers/Admin/InstallmentController.php",
            "app/Models/Order.php; app/Models/OrderDetail.php; app/Models/CouponFlashSale.php",
            "resources/views/frontend/cart/*.blade.php; resources/views/admin/installments/*.blade.php",
        ],
        "result": "Hoàn thành luồng mua hàng từ giỏ hàng đến thanh toán, áp dụng ưu đãi, vận chuyển, in bill và trả góp.",
    },
    {
        "stt": "3",
        "member": "Nguyễn Thanh Hiền\nPhân hệ chi tiết sản phẩm, hồ sơ, hậu mãi",
        "features": [
            "Trang chi tiết sản phẩm, biến thể, thông số, combo và cross-sell.",
            "Đánh giá, bình luận, phản hồi phân cấp và báo cáo vi phạm.",
            "Danh sách yêu thích và quản lý wishlist trong hồ sơ.",
            "Profile: cập nhật thông tin, đổi mật khẩu, địa chỉ giao hàng.",
            "Dashboard quản trị và thống kê vận hành.",
            "Sổ quỹ thu chi: CRUD, lọc và xóa hàng loạt.",
            "Chính sách bảo hành/đổi trả và tra cứu bảo hành.",
            "Video review: xem, thích, bình luận, báo cáo.",
            "Đa ngôn ngữ và đồng bộ bản dịch dữ liệu.",
            "Phiếu sửa chữa, hóa đơn dịch vụ và quản lý trả góp admin.",
        ],
        "files": [
            "app/Http/Controllers/Frontend/ProductController.php; resources/views/frontend/products/show.blade.php",
            "app/Http/Controllers/ReviewController.php; app/Http/Controllers/WishlistController.php",
            "app/Http/Controllers/ProfileController.php; resources/views/frontend/profile.blade.php",
            "app/Http/Controllers/Admin/DashboardController.php; app/Http/Controllers/CashbookController.php",
            "app/Http/Controllers/VideoController.php; app/Http/Controllers/Admin/VideoManagementController.php",
            "app/Http/Controllers/Frontend/WarrantyController.php; app/Http/Controllers/Admin/RepairTicketInvoiceController.php",
            "app/Http/Controllers/Admin/ServiceInvoiceController.php; app/Services/TranslationService.php",
        ],
        "result": "Hoàn thành trải nghiệm chi tiết sản phẩm, hồ sơ khách hàng, đánh giá, video, bảo hành và nghiệp vụ hậu mãi.",
    },
    {
        "stt": "4",
        "member": "Văn Nguyễn Xuân Hòa\nPhân hệ xác thực, phân quyền, bảo mật",
        "features": [
            "Đăng nhập/đăng ký, validate dữ liệu và phân hướng theo vai trò.",
            "2FA OTP, bật/tắt xác thực hai lớp và quản lý phiên đăng nhập.",
            "Khôi phục mật khẩu bằng OTP/email và đặt mật khẩu mới.",
            "Google SSO và liên kết tài khoản OAuth.",
            "Phân quyền RBAC theo Admin/Manager/Staff.",
            "Quản lý lịch sử đăng nhập, thiết bị và thu hồi session.",
            "CRUD nhân viên, đổi trạng thái, thao tác hàng loạt, xuất Excel/PDF.",
            "KPI nhân viên và xem chi tiết hiệu suất.",
            "Nhật ký hoạt động/Audit log, masking dữ liệu và xác minh hash.",
            "Tùy biến theme admin và trung tâm chat nội bộ.",
        ],
        "files": [
            "app/Http/Controllers/Auth/AuthController.php; app/Http/Controllers/Auth/TwoFactorController.php",
            "app/Http/Controllers/Auth/ForgotPasswordController.php; app/Http/Controllers/Auth/SocialController.php",
            "app/Http/Controllers/Admin/UserController.php; app/Http/Controllers/Admin/RoleController.php",
            "app/Http/Controllers/Admin/EmployeeController.php; app/Http/Controllers/Admin/KPIController.php",
            "app/Http/Controllers/Admin/ActivityLogController.php; app/Traits/HasAuditLog.php",
            "app/Http/Controllers/Admin/ThemeSettingController.php; app/Http/Controllers/Admin/ChatController.php",
            "app/Http/Middleware/CheckRole.php; app/Http/Middleware/IsAdmin.php",
        ],
        "result": "Hoàn thành nền tảng bảo mật tài khoản, phân quyền, quản trị nhân viên, KPI và giám sát hoạt động hệ thống.",
    },
    {
        "stt": "5",
        "member": "Đặng Đăng Nguyên\nPhân hệ sản phẩm, danh mục, kho hàng",
        "features": [
            "CRUD danh mục đa cấp và bản dịch danh mục.",
            "CRUD nhà cung cấp và thông tin liên hệ.",
            "CRUD sản phẩm, hình ảnh, giá, trạng thái và SEO.",
            "Biến thể, thuộc tính, cấu hình thông số và dữ liệu lọc.",
            "Phiếu nhập kho, nhập biến thể và định danh IMEI/Serial.",
            "Đồng bộ tồn kho khi bán, nhập, hoàn hoặc đổi trạng thái đơn.",
            "Cảnh báo tồn kho an toàn và điều chuyển kho nội bộ.",
            "Kiểm kê kho, ghi nhận chênh lệch và cân bằng tồn kho.",
            "Lịch sử biến động kho theo sản phẩm/kho/thời gian.",
            "Import/Export sản phẩm bằng Excel và kiểm tra dữ liệu nhập.",
        ],
        "files": [
            "app/Http/Controllers/Admin/CategoryController.php; app/Http/Controllers/Admin/SupplierController.php",
            "app/Http/Controllers/Admin/ProductController.php; app/Http/Controllers/Admin/AttributeController.php",
            "app/Http/Controllers/Admin/PurchaseOrderController.php; app/Http/Controllers/Admin/InventoryController.php",
            "app/Http/Controllers/Admin/InventoryMovementController.php; app/Http/Controllers/Admin/InventoryAuditController.php",
            "app/Http/Controllers/Admin/WarehouseTransferController.php; app/Services/InventoryService.php",
            "app/Imports/ProductImport.php; app/Exports/ProductExport.php",
            "app/Models/Product.php; app/Models/ProductVariant.php; app/Models/InventoryItem.php; app/Models/WarehouseTransfer.php",
        ],
        "result": "Hoàn thành lõi quản lý hàng hóa: danh mục, sản phẩm, biến thể, nhập kho, tồn kho, kiểm kê và điều chuyển kho.",
    },
]


def set_run_font(run, size=9, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, value, size=8.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(value)
    set_run_font(run, size=size, bold=bold)


def add_lines(cell, lines, numbered=True, size=8):
    cell.text = ""
    for idx, line in enumerate(lines, start=1):
        p = cell.paragraphs[0] if idx == 1 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.05
        prefix = f"{idx}. " if numbered else "- "
        run = p.add_run(prefix + line)
        set_run_font(run, size=size)


def set_width(cell, width_inches):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def delete_between(heading_para, next_para):
    parent = heading_para._element.getparent()
    children = list(parent)
    start = children.index(heading_para._element)
    end = children.index(next_para._element)
    for element in children[start + 1:end]:
        parent.remove(element)


def move_after(anchor, elements):
    current = anchor
    for element in elements:
        old_parent = element.getparent()
        if old_parent is not None:
            old_parent.remove(element)
        current.addnext(element)
        current = element


def main():
    doc = Document(INPUT_PATH)

    # Keep the detailed report table readable. This section-wide landscape setting
    # is acceptable for the large source-mapping table and avoids cramped columns.
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    heading_idx = None
    next_idx = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == "1.2. Bảng phân chia chức năng và file mã nguồn chi tiết theo thành viên":
            heading_idx = i
        elif heading_idx is not None and text == "CHƯƠNG 2: LÝ DO CHỌN ĐỀ TÀI VÀ MÔ TẢ NGHIỆP VỤ":
            next_idx = i
            break

    if heading_idx is None or next_idx is None:
        raise RuntimeError("Không tìm thấy đúng vùng nội dung 1.2 cần chỉnh.")

    heading_para = doc.paragraphs[heading_idx]
    next_para = doc.paragraphs[next_idx]
    delete_between(heading_para, next_para)

    intro = doc.add_paragraph()
    intro.style = doc.styles["Normal"]
    intro.paragraph_format.space_after = Pt(6)
    intro.paragraph_format.line_spacing = 1.15
    run = intro.add_run(
        "Bảng dưới đây tổng hợp phạm vi chức năng và mã nguồn trọng tâm theo từng thành viên. "
        "Nội dung được gom theo phân hệ để dễ đối chiếu khi chấm báo cáo, tránh liệt kê rời rạc theo tên nhánh Git."
    )
    set_run_font(run, size=10)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    headers = ["STT", "Thành viên / Phân hệ", "10 chức năng chính", "File mã nguồn trọng tâm", "Kết quả bàn giao"]
    widths = [0.45, 1.85, 4.05, 4.15, 2.0]

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "D9EAF7")
        set_cell_text(cell, header, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_width(cell, widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for item in ROWS:
        row = table.add_row()
        cells = row.cells
        for idx, width in enumerate(widths):
            set_width(cells[idx], width)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        set_cell_text(cells[0], item["stt"], size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], item["member"], size=8.5, bold=True)
        add_lines(cells[2], item["features"], numbered=True, size=7.7)
        add_lines(cells[3], item["files"], numbered=False, size=7.5)
        set_cell_text(cells[4], item["result"], size=8)

    move_after(heading_para._element, [intro._element, table._element])
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
