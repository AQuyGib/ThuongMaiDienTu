from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX_PATH = Path(r"G:\ThuongMaiDienTu\Nhom_g_bao_cao_lap_trinh_bach_end_web2_chinh_sua_Xuan_Hoa.docx")
OUTPUT_PATH = Path(r"G:\ThuongMaiDienTu\Nhom_g_bao_cao_lap_trinh_bach_end_web2_chinh_sua_Xuan_Hoa_bang_tong_quan_da_chinh.docx")


OVERVIEW_TABLES = [
    [
        ("1", "AI Chatbot (Gemini RAG & đặt lịch): tư vấn sản phẩm theo dữ liệu thực, lưu lịch sử hội thoại và hỗ trợ tạo phiếu sửa chữa từ nội dung chat.", "Rất khó", "2.0 đ", "100%"),
        ("2", "CRUD bài viết công nghệ: tạo, sửa, xóa, duyệt và hiển thị bài viết/blog UGC trên frontend và trang quản trị.", "Khó", "1.0 đ", "100%"),
        ("3", "CRUD khách hàng (CRM): quản lý hồ sơ khách hàng, tìm kiếm, phân loại, khóa/mở và theo dõi lịch sử tương tác.", "Trung bình", "1.0 đ", "100%"),
        ("4", "Flash Sale giờ vàng: cấu hình khung giờ, sản phẩm khuyến mãi, giới hạn số lượng và chống bán vượt tồn kho.", "Khó", "1.5 đ", "100%"),
        ("5", "Gợi ý bán chéo & combo: đề xuất sản phẩm mua kèm, combo bundle và gợi ý theo thương hiệu/flash sale/hành vi giỏ hàng.", "Khó", "1.0 đ", "100%"),
        ("6", "Bộ lọc nâng cao: lọc sản phẩm theo danh mục, thương hiệu, khoảng giá và thông số kỹ thuật động.", "Trung bình", "1.0 đ", "100%"),
        ("7", "So sánh sản phẩm: đồng bộ danh sách so sánh, giới hạn số lượng và hiển thị bảng thông số khác biệt.", "Trung bình", "0.75 đ", "100%"),
        ("8", "Tích điểm thành viên: tính điểm theo đơn hàng, ví điểm tiêu dùng, hạng thành viên và lịch sử đổi thưởng.", "Trung bình", "0.75 đ", "100%"),
        ("9", "Thông báo hệ thống: tạo chiến dịch thông báo, lọc người nhận, đánh dấu đã đọc và cảnh báo tồn kho thấp.", "Trung bình", "0.75 đ", "100%"),
        ("10", "Trang chủ & khối hiển thị sản phẩm: quản lý banner/section, sắp xếp khối trang chủ và hiển thị sản phẩm nổi bật.", "Dễ", "0.25 đ", "100%"),
    ],
    [
        ("1", "Quản lý giỏ hàng: thêm, cập nhật, xóa, chọn sản phẩm, đồng bộ số lượng và kiểm tra tồn kho khi đặt hàng.", "Khó", "1.5 đ", "100%"),
        ("2", "Thanh toán đơn hàng: xử lý form checkout, tạo đơn, áp dụng điểm/voucher và cập nhật trạng thái thanh toán.", "Rất khó", "1.5 đ", "100%"),
        ("3", "Quét mã QR thanh toán: sinh QR chuyển khoản cho đơn hàng, hỗ trợ xác nhận nhanh tại bước thanh toán.", "Khó", "1.0 đ", "100%"),
        ("4", "Tính phí vận chuyển: tính phí giao hàng theo thông tin nhận hàng, khu vực và dữ liệu đơn hàng.", "Trung bình", "1.0 đ", "100%"),
        ("5", "Mã giảm giá: kiểm tra điều kiện voucher, giới hạn lượt dùng, loại giảm theo phần trăm/số tiền và áp dụng vào đơn.", "Khó", "1.0 đ", "100%"),
        ("6", "Tra cứu đơn hàng: tìm kiếm đơn, xem trạng thái xử lý, thông tin giao hàng và lịch sử cập nhật.", "Trung bình", "1.0 đ", "100%"),
        ("7", "In hóa đơn đơn hàng: xuất giao diện bill/invoice cho đơn mua hàng, phục vụ in trực tiếp khi hoàn tất giao dịch.", "Trung bình", "0.75 đ", "100%"),
        ("8", "Quản lý trả góp phía khách hàng: tiếp nhận hồ sơ trả góp, tính kỳ hạn và gửi thông báo xử lý hồ sơ.", "Trung bình", "0.75 đ", "100%"),
        ("9", "Hủy đơn & hoàn/đổi trả: xử lý hủy đơn, timeout đơn, trả lại tồn kho và cập nhật trạng thái nghiệp vụ.", "Trung bình", "0.75 đ", "100%"),
        ("10", "Bán hàng tại quầy/POS cơ bản: hỗ trợ luồng giỏ hàng, thanh toán và in bill cho nhân viên bán trực tiếp.", "Dễ", "0.75 đ", "100%"),
    ],
    [
        ("1", "Trang chi tiết sản phẩm: hiển thị biến thể, thông số kỹ thuật, tồn kho, đánh giá, combo và sản phẩm mua kèm.", "Khó", "1.5 đ", "100%"),
        ("2", "Đánh giá & bình luận: thêm/sửa/xóa/report review, phản hồi phân cấp và kiểm duyệt nội dung vi phạm.", "Khó", "1.25 đ", "100%"),
        ("3", "Danh sách yêu thích: lưu/xóa sản phẩm yêu thích, xóa toàn bộ và hiển thị trong hồ sơ người dùng.", "Trung bình", "0.75 đ", "100%"),
        ("4", "Hồ sơ cá nhân: cập nhật thông tin, đổi mật khẩu, quản lý ảnh đại diện và sổ địa chỉ giao nhận.", "Trung bình", "1.0 đ", "100%"),
        ("5", "Dashboard quản trị: tổng hợp doanh thu, đơn hàng, khách hàng, tồn kho và KPI vận hành.", "Rất khó", "1.5 đ", "100%"),
        ("6", "Sổ quỹ thu chi: quản lý khoản thu/chi, lọc dữ liệu, cập nhật và xóa nhiều dòng giao dịch.", "Khó", "1.0 đ", "100%"),
        ("7", "Chính sách bảo hành & đổi trả: xây dựng trang chính sách, tra cứu bảo hành và tiếp nhận yêu cầu bảo hành.", "Trung bình", "0.75 đ", "100%"),
        ("8", "Video review & bình luận video: đăng/xem video, lượt thích, lượt xem, bình luận và báo cáo vi phạm.", "Trung bình", "0.75 đ", "100%"),
        ("9", "Đa ngôn ngữ: chuyển locale, đồng bộ bản dịch danh mục/sản phẩm/trang và hiển thị nội dung theo ngôn ngữ.", "Trung bình", "0.75 đ", "100%"),
        ("10", "Phiếu sửa chữa, dịch vụ & trả góp admin: quản lý phiếu sửa chữa, hóa đơn dịch vụ và phê duyệt hồ sơ trả góp.", "Khó", "0.75 đ", "100%"),
    ],
    [
        ("1", "Đăng nhập/đăng ký: xác thực tài khoản, validate dữ liệu, mã hóa mật khẩu và phân hướng theo vai trò.", "Trung bình", "1.0 đ", "100%"),
        ("2", "Xác thực 2FA OTP: gửi, xác minh, bật/tắt bảo mật hai lớp và quản lý phiên đăng nhập an toàn.", "Khó", "1.25 đ", "100%"),
        ("3", "Khôi phục mật khẩu: gửi OTP/email, xác minh mã và cập nhật mật khẩu mới theo quy trình bảo mật.", "Trung bình", "1.0 đ", "100%"),
        ("4", "Google SSO: đăng nhập bằng Google, liên kết/tạo tài khoản và lưu thông tin nhận dạng OAuth.", "Khó", "1.0 đ", "100%"),
        ("5", "Phân quyền RBAC: quản lý vai trò, quyền truy cập admin/manager/staff và middleware kiểm tra role.", "Rất khó", "1.5 đ", "100%"),
        ("6", "Quản lý lịch sử đăng nhập: ghi nhận IP, thiết bị, phiên hoạt động và cho phép thu hồi session.", "Trung bình", "1.0 đ", "100%"),
        ("7", "CRUD nhân viên: tạo, cập nhật, xóa, đổi trạng thái, thao tác hàng loạt và xuất Excel/PDF.", "Khó", "1.0 đ", "100%"),
        ("8", "KPI nhân viên: tổng hợp chỉ số doanh số/đơn hàng, xem chi tiết theo nhân viên và hỗ trợ đánh giá hiệu suất.", "Trung bình", "0.75 đ", "100%"),
        ("9", "Nhật ký hoạt động/Audit log: ghi vết hành động nhạy cảm, che dữ liệu riêng tư và xác minh hash log.", "Khó", "1.0 đ", "100%"),
        ("10", "Themes & trung tâm giao tiếp nội bộ: tùy biến giao diện admin và quản lý phòng chat/thành viên/tin nhắn.", "Trung bình", "0.5 đ", "100%"),
    ],
    [
        ("1", "CRUD danh mục: quản lý cây danh mục sản phẩm, slug, ảnh, trạng thái và bản dịch danh mục.", "Trung bình", "1.0 đ", "100%"),
        ("2", "CRUD nhà cung cấp: quản lý thông tin liên hệ, tìm kiếm nhà cung cấp và liên kết lịch sử nhập hàng.", "Trung bình", "0.75 đ", "100%"),
        ("3", "CRUD sản phẩm: tạo/sửa/xóa sản phẩm, ảnh, thương hiệu, giá, tồn kho, trạng thái và dữ liệu SEO.", "Khó", "1.25 đ", "100%"),
        ("4", "Biến thể & thuộc tính sản phẩm: quản lý RAM/màu/CPU/GPU, thông số kỹ thuật và dữ liệu lọc theo danh mục.", "Khó", "1.0 đ", "100%"),
        ("5", "Phiếu nhập kho & định danh IMEI/Serial: lập PO, nhập biến thể, sinh item tồn kho và kiểm soát mã định danh.", "Rất khó", "1.5 đ", "100%"),
        ("6", "Đồng bộ tồn kho: tự động trừ/hoàn tồn theo trạng thái đơn hàng, flash sale và thao tác kho.", "Khó", "1.0 đ", "100%"),
        ("7", "Cảnh báo tồn kho & điều chuyển kho: phát hiện dưới ngưỡng an toàn, tạo phiếu chuyển kho và hoàn tất/hủy điều chuyển.", "Khó", "1.0 đ", "100%"),
        ("8", "Kiểm kê và cân bằng kho: lập phiếu kiểm kê, ghi nhận chênh lệch thực tế và reconcile số lượng.", "Trung bình", "0.75 đ", "100%"),
        ("9", "Lịch sử biến động kho: ghi log nhập/xuất/bán/hoàn/trả kho và lọc theo sản phẩm, kho, thời gian.", "Trung bình", "0.75 đ", "100%"),
        ("10", "Import/Export sản phẩm: tải mẫu Excel, nhập hàng loạt, xuất danh sách sản phẩm và kiểm tra dữ liệu trước khi lưu.", "Trung bình", "1.0 đ", "100%"),
    ],
]


def set_cell(cell, text, bold=False, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)


def main():
    doc = Document(DOCX_PATH)

    # Fix the inconsistent score line for Huynh Van Vinh Em.
    for paragraph in doc.paragraphs:
        if "Tổng điểm cá nhân đạt được: 8 / 10.0 điểm." in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "Tổng điểm cá nhân đạt được: 8 / 10.0 điểm.",
                "Tổng điểm cá nhân đạt được: 10.0 / 10.0 điểm.",
            )

    for table_index, rows in enumerate(OVERVIEW_TABLES):
        table = doc.tables[table_index]
        for row_offset, row_values in enumerate(rows, start=1):
            for col_index, value in enumerate(row_values):
                align = WD_ALIGN_PARAGRAPH.CENTER if col_index in (0, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
                set_cell(table.rows[row_offset].cells[col_index], value, align=align)

        total_row = table.rows[11]
        set_cell(total_row.cells[0], "Tổng", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(total_row.cells[1], "10 chức năng chính", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(total_row.cells[2], "", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(total_row.cells[3], "10.0 đ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(total_row.cells[4], "Đạt", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
